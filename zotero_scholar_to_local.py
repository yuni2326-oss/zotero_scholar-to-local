from __future__ import annotations

import argparse
import asyncio
import configparser
import datetime
import html
import json
import random
import re
import sqlite3
import string
import sys
import threading
import urllib.parse
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Optional

# ── docx 의존성 (optional) ───────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

FONT = "Malgun Gothic"


def _safe_filename(keyword: str) -> str:
    """파일명에 사용할 수 없는 문자를 _로 치환한다."""
    return re.sub(r'[\\/:*?"<>|]', "_", keyword)


# ── 데이터 클래스 ─────────────────────────────────────────────────────────────

@dataclass
class ZoteroPaths:
    profiles_ini: Path
    profile_dir: Path
    data_dir: Path
    db_path: Path


@dataclass
class ScholarPaper:
    title: str
    url: str
    authors: list[str]
    year: Optional[str]
    venue: Optional[str]
    abstract: Optional[str] = None   # OpenAlex 등에서 채움
    doi: Optional[str] = None        # OpenAlex에서 채움
    source: str = "scholar"          # "scholar" | "openalex" | "semantic_scholar"


# ── Zotero 경로 탐색 ──────────────────────────────────────────────────────────

def parse_user_pref(prefs_text: str, key: str) -> Optional[str]:
    pattern = rf'user_pref\("{re.escape(key)}",\s*(.+?)\);'
    match = re.search(pattern, prefs_text)
    if not match:
        return None
    raw = match.group(1).strip()
    if raw.lower() in {"true", "false"}:
        return raw.lower()
    if raw.startswith('"') and raw.endswith('"'):
        value = raw[1:-1]
        return bytes(value, "utf-8").decode("unicode_escape")
    return raw


def resolve_zotero_paths() -> ZoteroPaths:
    appdata = Path.home() / "AppData" / "Roaming"
    profiles_ini = appdata / "Zotero" / "Zotero" / "profiles.ini"
    if not profiles_ini.exists():
        raise FileNotFoundError(f"profiles.ini not found: {profiles_ini}")

    cfg = configparser.RawConfigParser()
    cfg.read(profiles_ini, encoding="utf-8")

    default_section = None
    for section in cfg.sections():
        if section.startswith("Profile") and cfg.get(section, "Default", fallback="0") == "1":
            default_section = section
            break
    if default_section is None:
        raise RuntimeError("No default Zotero profile found in profiles.ini")

    is_relative = cfg.get(default_section, "IsRelative", fallback="1") == "1"
    profile_path = cfg.get(default_section, "Path", fallback="")
    if not profile_path:
        raise RuntimeError("Default profile has no path")

    profile_dir = profiles_ini.parent / profile_path if is_relative else Path(profile_path)

    prefs_js = profile_dir / "prefs.js"
    if not prefs_js.exists():
        raise FileNotFoundError(f"prefs.js not found: {prefs_js}")

    prefs_text = prefs_js.read_text(encoding="utf-8", errors="ignore")
    use_data_dir = parse_user_pref(prefs_text, "extensions.zotero.useDataDir")
    custom_data_dir = parse_user_pref(prefs_text, "extensions.zotero.dataDir")

    data_dir = Path(custom_data_dir) if use_data_dir == "true" and custom_data_dir else profile_dir
    db_path = data_dir / "zotero.sqlite"
    if not db_path.exists():
        raise FileNotFoundError(f"zotero.sqlite not found: {db_path}")

    return ZoteroPaths(profiles_ini=profiles_ini, profile_dir=profile_dir,
                       data_dir=data_dir, db_path=db_path)


# ── 번역 ──────────────────────────────────────────────────────────────────────

def is_mostly_english(text: str) -> bool:
    ascii_letters = sum(ch in string.ascii_letters for ch in text)
    return ascii_letters >= max(1, len(text) // 3)


def translate_to_english(keyword: str) -> str:
    if is_mostly_english(keyword):
        return keyword
    params = {"client": "gtx", "sl": "auto", "tl": "en", "dt": "t", "q": keyword}
    url = "https://translate.googleapis.com/translate_a/single?" + urllib.parse.urlencode(params)
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req, timeout=20) as resp:
        data = json.loads(resp.read().decode("utf-8"))
    translated = "".join(seg[0] for seg in data[0] if seg and seg[0]).strip()
    return translated or keyword


def translate_to_korean(text: str, max_chars: int = 4500) -> str:
    """영문 텍스트를 한국어로 번역 (Google Translate 무료 API). 실패 시 원문 반환."""
    if not text:
        return text
    text = text[:max_chars]  # API 글자 수 제한 대비
    params = {"client": "gtx", "sl": "auto", "tl": "ko", "dt": "t", "q": text}
    url = "https://translate.googleapis.com/translate_a/single?" + urllib.parse.urlencode(params)
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        translated = "".join(seg[0] for seg in data[0] if seg and seg[0]).strip()
        return translated or text
    except Exception:
        return text  # 번역 실패 시 원문 유지


# ── Google Scholar 검색 ───────────────────────────────────────────────────────

def clean_html_text(raw: str) -> str:
    return html.unescape(re.sub(r"<[^>]+>", "", raw)).strip()


def parse_scholar_result_blocks(page_html: str) -> list[ScholarPaper]:
    blocks = re.findall(r'<div class="gs_r gs_or gs_scl".*?</div>\s*</div>', page_html, flags=re.S)
    papers: list[ScholarPaper] = []
    for block in blocks:
        title_match = re.search(r'<h3 class="gs_rt"[^>]*>(.*?)</h3>', block, flags=re.S)
        if not title_match:
            continue
        h3 = title_match.group(1)
        link_match = re.search(r'<a href="([^"]+)"[^>]*>(.*?)</a>', h3, flags=re.S)
        if link_match:
            url, title = html.unescape(link_match.group(1)), clean_html_text(link_match.group(2))
        else:
            url, title = "", clean_html_text(h3)
        if not title or title.lower().startswith("citation"):
            continue
        meta_match = re.search(r'<div class="gs_a"[^>]*>(.*?)</div>', block, flags=re.S)
        meta = clean_html_text(meta_match.group(1)) if meta_match else ""
        year_match = re.search(r'(19|20)\d{2}', meta)
        year = year_match.group(0) if year_match else None
        parts = [p.strip() for p in meta.split(" - ") if p.strip()]
        authors_part = parts[0] if parts else ""
        venue = parts[1] if len(parts) > 1 else None
        authors = [a.strip() for a in authors_part.split(",") if a.strip() and a.strip() != "…"]
        papers.append(ScholarPaper(title=title, url=url, authors=authors,
                                   year=year, venue=venue, source="scholar"))
    return papers


def search_google_scholar_recent(english_query: str, limit: int,
                                 years_back: Optional[int] = None) -> list[ScholarPaper]:
    params = {"hl": "en", "q": english_query, "as_sdt": "0,5", "scisbd": "1"}
    if years_back:
        params["as_ylo"] = str(datetime.date.today().year - years_back)
    url = "https://scholar.google.com/scholar?" + urllib.parse.urlencode(params)
    req = urllib.request.Request(url, headers={
        "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                       "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36")
    })
    with urllib.request.urlopen(req, timeout=20) as resp:
        page_html = resp.read().decode("utf-8", errors="ignore")
    return parse_scholar_result_blocks(page_html)[:limit]


# ── OpenAlex 검색 ─────────────────────────────────────────────────────────────

def _reconstruct_abstract(inverted_index: Optional[dict]) -> Optional[str]:
    """OpenAlex inverted_index 포맷 → 평문 abstract 복원"""
    if not inverted_index:
        return None
    words: list[tuple[int, str]] = []
    for word, positions in inverted_index.items():
        for pos in positions:
            words.append((pos, word))
    words.sort()
    return " ".join(w for _, w in words)


def search_openalex(english_query: str, limit: int,
                    years_back: Optional[int] = None) -> list[ScholarPaper]:
    params = {
        "search": english_query,
        "sort": "publication_date:desc",
        "per-page": str(limit),
        "mailto": "research@example.com",  # OpenAlex polite pool
    }
    if years_back:
        from_year = datetime.date.today().year - years_back
        params["filter"] = f"publication_year:>{from_year - 1}"
    url = "https://api.openalex.org/works?" + urllib.parse.urlencode(params)
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    try:
        with urllib.request.urlopen(req, timeout=20) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except Exception as exc:
        print(f"[WARN] OpenAlex 검색 실패: {exc}")
        return []

    papers: list[ScholarPaper] = []
    for work in data.get("results", []):
        title = work.get("title") or ""
        if not title:
            continue

        # 저자
        authors = [
            a.get("author", {}).get("display_name", "")
            for a in work.get("authorships", [])[:8]
            if a.get("author", {}).get("display_name")
        ]

        # 연도
        pub_date = work.get("publication_date") or ""
        year = pub_date[:4] if pub_date else None

        # 저널
        venue = None
        loc = work.get("primary_location") or {}
        source = loc.get("source") or {}
        if source:
            venue = source.get("display_name")

        # DOI / URL
        doi = work.get("doi")
        if doi and doi.startswith("https://doi.org/"):
            doi = doi[len("https://doi.org/"):]
        url_val = work.get("doi") or ""

        # Abstract
        abstract = _reconstruct_abstract(work.get("abstract_inverted_index"))

        papers.append(ScholarPaper(
            title=title, url=url_val, authors=authors,
            year=year, venue=venue, abstract=abstract,
            doi=doi, source="openalex",
        ))

    return papers


# ── Semantic Scholar 검색 ──────────────────────────────────────────────────────

def search_semantic_scholar(english_query: str, limit: int,
                             years_back: Optional[int] = None) -> list[ScholarPaper]:
    import time as _time
    import urllib.error as _uerr

    params: dict[str, str] = {
        "query": english_query,
        "limit": str(limit),
        "fields": "title,authors,year,venue,externalIds,abstract",
    }
    if years_back:
        from_year = datetime.date.today().year - years_back
        params["year"] = f"{from_year}-{datetime.date.today().year}"

    url = "https://api.semanticscholar.org/graph/v1/paper/search?" + urllib.parse.urlencode(params)
    req = urllib.request.Request(url, headers={
        "User-Agent": "Mozilla/5.0 (compatible; research-tool/1.0)",
        "Accept": "application/json",
    })
    def _parse_retry_time(exc: "_uerr.HTTPError") -> str:
        """429 응답의 Date 헤더를 파싱해 재시도 가능 시각(로컬 시간)을 반환한다."""
        import email.utils as _eu
        date_str = exc.headers.get("Date", "")
        try:
            server_dt = _eu.parsedate_to_datetime(date_str)
            retry_dt = server_dt + datetime.timedelta(minutes=5)
            # UTC → 로컬 시간으로 변환
            retry_local = retry_dt.astimezone().strftime("%H:%M:%S")
            now_local = _time.strftime("%H:%M:%S")
            wait_sec = max(0, int((retry_dt.timestamp() - _time.time())))
            return f"{retry_local} 이후 재시도 가능 (현재 {now_local}, 약 {wait_sec}초 후)"
        except Exception:
            return "약 5분 후 재시도하세요."

    data = None
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except _uerr.HTTPError as exc:
        if exc.code == 429:
            retry_msg = _parse_retry_time(exc)
            print(f"[WARN] Semantic Scholar 요청 제한(429) -- {retry_msg}")
            return []
        print(f"[WARN] Semantic Scholar 검색 실패: {exc}")
        return []
    except Exception as exc:
        print(f"[WARN] Semantic Scholar 검색 실패: {exc}")
        return []
    if data is None:
        return []

    papers: list[ScholarPaper] = []
    for work in data.get("data", []):
        title = work.get("title") or ""
        if not title:
            continue

        authors = [
            a.get("name", "")
            for a in work.get("authors", [])[:8]
            if a.get("name")
        ]
        year = str(work["year"]) if work.get("year") else None
        venue = work.get("venue") or None
        abstract = work.get("abstract") or None

        ext_ids = work.get("externalIds") or {}
        doi = ext_ids.get("DOI") or None
        url_val = f"https://doi.org/{doi}" if doi else ""

        papers.append(ScholarPaper(
            title=title, url=url_val, authors=authors,
            year=year, venue=venue, abstract=abstract,
            doi=doi, source="semantic_scholar",
        ))

    return papers


# ── 병합 및 중복 제거 ──────────────────────────────────────────────────────────

def merge_and_deduplicate(scholar_papers: list[ScholarPaper],
                          openalex_papers: list[ScholarPaper],
                          semantic_scholar_papers: list[ScholarPaper] | None = None) -> list[ScholarPaper]:
    """Scholar 우선, 제목 소문자 기준 중복 제거. abstract/doi를 후순위 소스로 보완."""
    seen: dict[str, ScholarPaper] = {}

    for p in scholar_papers:
        key = re.sub(r'\s+', ' ', p.title.lower().strip())
        seen[key] = p

    for p in openalex_papers:
        key = re.sub(r'\s+', ' ', p.title.lower().strip())
        if key not in seen:
            seen[key] = p
        else:
            existing = seen[key]
            if not existing.abstract and p.abstract:
                existing.abstract = p.abstract
            if not existing.doi and p.doi:
                existing.doi = p.doi

    for p in (semantic_scholar_papers or []):
        key = re.sub(r'\s+', ' ', p.title.lower().strip())
        if key not in seen:
            seen[key] = p
        else:
            existing = seen[key]
            if not existing.abstract and p.abstract:
                existing.abstract = p.abstract
            if not existing.doi and p.doi:
                existing.doi = p.doi

    return list(seen.values())


# ── Zotero SQLite 쓰기 ────────────────────────────────────────────────────────

def make_zotero_key(existing_check) -> str:
    alphabet = "23456789ABCDEFGHIJKLMNPQRSTUVWXYZ"
    for _ in range(200):
        key = "".join(random.choice(alphabet) for _ in range(8))
        if not existing_check(key):
            return key
    raise RuntimeError("Could not generate a unique Zotero key")


def get_or_create_value_id(cur: sqlite3.Cursor, value: str) -> int:
    found = cur.execute("SELECT valueID FROM itemDataValues WHERE value = ?", (value,)).fetchone()
    if found:
        return int(found[0])
    cur.execute("INSERT INTO itemDataValues(value) VALUES (?)", (value,))
    return int(cur.lastrowid)


def get_or_create_creator(cur: sqlite3.Cursor, full_name: str) -> int:
    full_name = full_name.strip()
    if not full_name:
        raise ValueError("author name is empty")
    if " " in full_name:
        first_name, last_name = full_name.rsplit(" ", 1)
        field_mode = 0
    else:
        first_name, last_name = "", full_name
        field_mode = 1
    found = cur.execute(
        "SELECT creatorID FROM creators WHERE firstName = ? AND lastName = ? AND fieldMode = ?",
        (first_name, last_name, field_mode),
    ).fetchone()
    if found:
        return int(found[0])
    cur.execute("INSERT INTO creators(firstName, lastName, fieldMode) VALUES (?, ?, ?)",
                (first_name, last_name, field_mode))
    return int(cur.lastrowid)


def get_or_create_collection(cur: sqlite3.Cursor, library_id: int, name: str) -> int:
    found = cur.execute(
        "SELECT collectionID FROM collections WHERE libraryID = ? AND collectionName = ? AND parentCollectionID IS NULL",
        (library_id, name),
    ).fetchone()
    if found:
        return int(found[0])

    def key_exists(k):
        return bool(cur.execute("SELECT 1 FROM collections WHERE libraryID = ? AND key = ?",
                                (library_id, k)).fetchone())

    key = make_zotero_key(key_exists)
    cur.execute(
        "INSERT INTO collections(collectionName, parentCollectionID, libraryID, key, version, synced) VALUES (?, NULL, ?, ?, 0, 0)",
        (name, library_id, key),
    )
    return int(cur.lastrowid)


def find_existing_item_id(cur: sqlite3.Cursor, title: str, doi: Optional[str],
                          library_id: int) -> Optional[int]:
    if doi:
        row = cur.execute(
            "SELECT i.itemID FROM items i JOIN itemData d ON d.itemID = i.itemID AND d.fieldID = 8 "
            "JOIN itemDataValues v ON v.valueID = d.valueID WHERE i.libraryID = ? AND v.value = ? LIMIT 1",
            (library_id, doi),
        ).fetchone()
        if row:
            return int(row[0])
    row = cur.execute(
        "SELECT i.itemID FROM items i JOIN itemData d ON d.itemID = i.itemID AND d.fieldID = 1 "
        "JOIN itemDataValues v ON v.valueID = d.valueID WHERE i.libraryID = ? AND v.value = ? LIMIT 1",
        (library_id, title),
    ).fetchone()
    return int(row[0]) if row else None


def attach_item_to_collection(cur: sqlite3.Cursor, collection_id: int, item_id: int) -> None:
    if cur.execute("SELECT 1 FROM collectionItems WHERE collectionID = ? AND itemID = ?",
                   (collection_id, item_id)).fetchone():
        return
    order_index = cur.execute(
        "SELECT COALESCE(MAX(orderIndex), -1) + 1 FROM collectionItems WHERE collectionID = ?",
        (collection_id,),
    ).fetchone()[0]
    cur.execute("INSERT INTO collectionItems(collectionID, itemID, orderIndex) VALUES (?, ?, ?)",
                (collection_id, item_id, int(order_index)))


def insert_paper_as_journal_article(cur: sqlite3.Cursor, library_id: int,
                                    paper: ScholarPaper) -> int:
    existing_id = find_existing_item_id(cur, paper.title, paper.doi, library_id)
    if existing_id:
        return existing_id

    def item_key_exists(k):
        return bool(cur.execute("SELECT 1 FROM items WHERE libraryID = ? AND key = ?",
                                (library_id, k)).fetchone())

    item_key = make_zotero_key(item_key_exists)
    cur.execute("INSERT INTO items(itemTypeID, libraryID, key, version, synced) VALUES (22, ?, ?, 0, 0)",
                (library_id, item_key))
    item_id = int(cur.lastrowid)

    # fieldID: 1=title, 6=date, 8=DOI, 10=url, 15=language, 41=publicationTitle, 27=abstractNote
    field_pairs: list[tuple[int, str]] = [(1, paper.title), (15, "en")]
    if paper.year:
        field_pairs.append((6, paper.year))
    if paper.venue:
        field_pairs.append((41, paper.venue))
    if paper.url:
        field_pairs.append((10, paper.url))
    if paper.doi:
        field_pairs.append((8, paper.doi))
    if paper.abstract:
        field_pairs.append((27, paper.abstract))

    for field_id, value in field_pairs:
        value_id = get_or_create_value_id(cur, value)
        cur.execute("INSERT OR REPLACE INTO itemData(itemID, fieldID, valueID) VALUES (?, ?, ?)",
                    (item_id, field_id, value_id))

    for idx, author in enumerate(paper.authors[:8]):
        creator_id = get_or_create_creator(cur, author)
        cur.execute(
            "INSERT INTO itemCreators(itemID, creatorID, creatorTypeID, orderIndex) VALUES (?, ?, 10, ?)",
            (item_id, creator_id, idx),
        )
    return item_id


def get_user_library_id(cur: sqlite3.Cursor) -> int:
    row = cur.execute("SELECT libraryID FROM libraries WHERE type='user' LIMIT 1").fetchone()
    if not row:
        raise RuntimeError("User library not found")
    return int(row[0])


# ── NotebookLM 통합 ───────────────────────────────────────────────────────────

def read_collection_papers_from_zotero(db_path: Path,
                                        collection_name: Optional[str] = None) -> str:
    """Zotero SQLite에서 논문 title + abstract를 텍스트로 반환 (최대 150편).

    collection_name 지정 시 해당 컬렉션만, None이면 전체 라이브러리에서 읽는다.
    """
    try:
        with sqlite3.connect(str(db_path)) as con:
            if collection_name:
                row = con.execute(
                    "SELECT collectionID FROM collections WHERE libraryID = 1 AND collectionName = ?",
                    (collection_name,)
                ).fetchone()
                if not row:
                    return ""
                coll_id = row[0]
                items = con.execute(
                    "SELECT ci.itemID FROM collectionItems ci "
                    "JOIN items i ON ci.itemID = i.itemID "
                    "WHERE ci.collectionID = ? LIMIT 150",
                    (coll_id,)
                ).fetchall()
            else:
                items = con.execute(
                    "SELECT i.itemID FROM items i "
                    "WHERE i.libraryID = 1 "
                    "AND i.itemTypeID NOT IN (1, 14) "
                    "LIMIT 150",
                ).fetchall()
            parts = []
            for n, (item_id,) in enumerate(items, 1):
                vals = {
                    fid: val
                    for fid, val in con.execute(
                        "SELECT id.fieldID, idv.value "
                        "FROM itemData id JOIN itemDataValues idv "
                        "ON id.valueID = idv.valueID "
                        "WHERE id.itemID = ? AND id.fieldID IN (1, 27)",
                        (item_id,)
                    ).fetchall()
                }
                title = vals.get(1, "")
                abstract = vals.get(27, "")
                if title:
                    parts.append(f"{n}. {title}\n{abstract}\n")
            return "\n".join(parts)
    except Exception as exc:
        print(f"[WARN] read_collection_papers_from_zotero 오류: {exc}")
        return ""


async def analyze_with_notebooklm(db_path: Path, keyword: str,
                                   proposal: str,
                                   log_fn: Callable[[str], None]) -> None:
    """NotebookLM에서 Zotero 컬렉션 논문으로 연구 제안 타당성을 평가한다."""
    try:
        from notebooklm import NotebookLMClient
    except ImportError:
        log_fn("[ERROR] notebooklm-py 미설치: pip install 'notebooklm-py[browser]'")
        return

    text = read_collection_papers_from_zotero(db_path, keyword)
    if not text.strip():
        log_fn(f"[WARN] Zotero 컬렉션 '{keyword}'에 논문이 없습니다.")
        return

    log_fn("[NLM] NotebookLM 연결 중...")
    try:
        async with await NotebookLMClient.from_storage() as client:
            nb = await client.notebooks.create(f"[분석] {keyword}")
            log_fn(f"[NLM] 노트북 생성: {nb.title}")
            try:
                await client.sources.add_text(nb.id, f"{keyword} 논문 목록", text)
                log_fn(f"[NLM] 소스 추가 완료 ({len(text)}자)")

                question = (
                    f"아래 연구 제안의 타당성을 논문들을 바탕으로 한국어로 평가해주세요.\n\n"
                    f"[연구 제안]\n{proposal}\n\n"
                    f"[평가 항목]\n"
                    f"1. 주요 연구 트렌드와의 관계\n"
                    f"2. 연구 공백(gap) 및 차별성\n"
                    f"3. 타당성 종합 평가"
                )
                log_fn("[NLM] 분석 중... (수십 초 소요)")
                result = await client.chat.ask(nb.id, question)
                log_fn("\n[NLM 분석 결과]\n" + result.answer)
            finally:
                try:
                    await client.notebooks.delete(nb.id)
                    log_fn("[NLM] 임시 노트북 삭제 완료")
                except Exception:
                    log_fn(f"[WARN] 노트북 삭제 실패 - NotebookLM에서 수동 삭제 필요: [분석] {keyword}")
    except Exception as exc:
        if "login" in str(exc).lower() or "auth" in str(exc).lower():
            log_fn("[ERROR] NotebookLM 로그인 필요: notebooklm login 실행 후 재시도")
        else:
            log_fn(f"[ERROR] NotebookLM 오류: {exc}")


def open_notebooklm_analysis(db_path: Path, keyword: str,
                              proposal: str,
                              log_fn: Callable[[str], None]) -> None:
    """analyze_with_notebooklm의 동기 래퍼 (GUI 백그라운드 스레드에서 호출)."""
    asyncio.run(analyze_with_notebooklm(db_path, keyword, proposal, log_fn))


# ── Claude Code 분석 연동 ──────────────────────────────────────────────────────

def save_analysis_request(keyword: str, proposal: str, output_dir: Path,
                           db_path: Optional[Path] = None) -> Path:
    """분석 요청 JSON 저장. 논문 데이터를 직접 포함하므로 Zotero 앱 불필요."""
    safe_kw = _safe_filename(keyword)
    json_path = output_dir / f"{safe_kw}_analysis.json"

    papers_text = ""
    if db_path:
        papers_text = read_collection_papers_from_zotero(db_path, keyword)

    instructions = (
        f"이 JSON 파일의 'papers' 필드에 포함된 논문 목록을 바탕으로, "
        f"아래 연구 제안의 타당성을 한국어로 평가해주세요. "
        f"(Zotero MCP는 사용하지 않아도 됩니다.)\n\n"
        f"[연구 제안]\n{proposal}\n\n"
        f"[평가 항목]\n"
        f"1. 주요 연구 트렌드와의 관계\n"
        f"2. 연구 공백(gap) 및 차별성\n"
        f"3. 타당성 종합 평가"
    )
    data = {
        "keyword": keyword,
        "collection": keyword.strip(),
        "proposal": proposal,
        "created_at": datetime.date.today().isoformat(),
        "papers": papers_text,
        "instructions": instructions,
    }
    json_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return json_path


def open_claude_terminal(json_path: Path,
                         log_fn: Callable[[str], None] = print) -> None:
    """새 CMD 창에서 Claude Code를 열고 JSON 분석 요청을 시작한다."""
    import subprocess
    import shutil

    if shutil.which("claude") is None:
        log_fn("[WARN] claude CLI를 찾을 수 없습니다.")
        log_fn(f"[INFO] 직접 Claude Code를 열고 파일을 참조하세요: {json_path}")
        return

    abs_path = str(json_path.resolve()).replace("\\", "/")
    prompt = f'Read "{abs_path}" and follow the instructions field'
    try:
        subprocess.Popen(
            ["cmd", "/k", "claude", prompt],
            creationflags=subprocess.CREATE_NEW_CONSOLE,
            cwd=str(json_path.parent),
        )
        log_fn(f"[INFO] Claude Code 터미널을 열었습니다.")
        log_fn(f"[INFO] 분석 파일: {json_path.name}")
    except Exception as exc:
        log_fn(f"[WARN] 터미널 실행 실패: {exc}")
        log_fn(f"[INFO] 수동으로 Claude Code를 열고 아래 파일을 참조하세요: {json_path}")


# ── docx 생성 ─────────────────────────────────────────────────────────────────

def _set_run_font(run, font_name: str = FONT) -> None:
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)


def generate_summary_docx(papers: list[ScholarPaper], keyword: str, output_path: Path,
                          log_fn: Callable[[str], None] = print) -> None:
    if not _DOCX_AVAILABLE:
        log_fn("[WARN] python-docx 미설치. docx 생성을 건너뜁니다. (pip install python-docx)")
        return

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)

    def add_run(para, text, bold=False, size=10, color=None):
        run = para.add_run(text)
        _set_run_font(run)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        return run

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"[{keyword}] 논문 요약 보고서", bold=True, size=18,
            color=RGBColor(0x1F, 0x49, 0x7D))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, f"작성일: {datetime.date.today().strftime('%Y년 %m월 %d일')}  |  "
                f"총 {len(papers)}편  |  출처: Google Scholar + OpenAlex + Semantic Scholar",
            size=9, color=RGBColor(0x70, 0x70, 0x70))
    doc.add_paragraph()

    # 소스별 분류
    scholar_papers = [p for p in papers if p.source == "scholar"]
    openalex_papers = [p for p in papers if p.source == "openalex"]
    ss_papers = [p for p in papers if p.source == "semantic_scholar"]

    def write_section(title: str, section_papers: list[ScholarPaper], start_num: int) -> int:
        if not section_papers:
            return start_num
        ph = doc.add_paragraph()
        add_run(ph, title, bold=True, size=13, color=RGBColor(0x2E, 0x74, 0xB5))

        for i, paper in enumerate(section_papers, start=start_num):
            pt = doc.add_paragraph()
            add_run(pt, f"{i}. {paper.title}", bold=True, size=11,
                    color=RGBColor(0x1F, 0x49, 0x7D))

            # 메타 정보
            pm = doc.add_paragraph()
            pm.paragraph_format.left_indent = Inches(0.2)
            add_run(pm, "저자: ", bold=True, size=9)
            add_run(pm, ("; ".join(paper.authors[:5]) + (" 외" if len(paper.authors) > 5 else ""))
                    if paper.authors else "N/A", size=9)
            add_run(pm, "  |  저널: ", bold=True, size=9)
            add_run(pm, paper.venue or "N/A", size=9)
            add_run(pm, "  |  연도: ", bold=True, size=9)
            add_run(pm, paper.year or "N/A", size=9)
            if paper.doi:
                add_run(pm, "  |  DOI: ", bold=True, size=9)
                add_run(pm, paper.doi, size=9, color=RGBColor(0x00, 0x70, 0xC0))

            # Abstract → 한국어 번역
            if paper.abstract:
                log_fn(f"  [{i}] 초록 번역 중...")
                ko_abstract = translate_to_korean(paper.abstract)
                pa = doc.add_paragraph()
                pa.paragraph_format.left_indent = Inches(0.2)
                add_run(pa, "초록: ", bold=True, size=10)
                add_run(pa, ko_abstract[:800] + ("…" if len(ko_abstract) > 800 else ""),
                        size=10)

            doc.add_paragraph()

        return start_num + len(section_papers)

    num = 1
    num = write_section("■ Google Scholar 검색 결과", scholar_papers, num)
    if openalex_papers:
        doc.add_page_break()
    num = write_section("■ OpenAlex 검색 결과", openalex_papers, num)
    if ss_papers:
        doc.add_page_break()
    num = write_section("■ Semantic Scholar 검색 결과", ss_papers, num)

    # 맺음말
    doc.add_page_break()
    pe = doc.add_paragraph()
    pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(pe, "── 요약 보고서 끝 ──", size=9, color=RGBColor(0xA0, 0xA0, 0xA0))
    pn = doc.add_paragraph()
    pn.paragraph_format.left_indent = Inches(0.2)
    add_run(pn, f"※ 키워드 '{keyword}'로 Google Scholar + OpenAlex + Semantic Scholar에서 검색한 결과입니다.",
            size=8, color=RGBColor(0x80, 0x80, 0x80))

    doc.save(str(output_path))
    log_fn(f"[DONE] docx 저장 완료: {output_path}")


# ── 메인 파이프라인 ───────────────────────────────────────────────────────────

def run_pipeline(db_path: Path, keyword: str, limit: int,
                 output_path: Optional[Path],
                 years_back: Optional[int] = None,
                 log_fn: Callable[[str], None] = print,
                 proposal: str = "") -> None:
    english_keyword = translate_to_english(keyword)
    log_fn(f"[INFO] 번역 키워드: {english_keyword}")

    # 검색
    log_fn("[INFO] Google Scholar 검색 중...")
    scholar_papers = search_google_scholar_recent(english_keyword, limit=limit,
                                                  years_back=years_back)
    log_fn(f"[INFO] Scholar 결과: {len(scholar_papers)}편")

    log_fn("[INFO] OpenAlex 검색 중...")
    openalex_papers = search_openalex(english_keyword, limit=limit, years_back=years_back)
    log_fn(f"[INFO] OpenAlex 결과: {len(openalex_papers)}편")

    log_fn("[INFO] Semantic Scholar 검색 중...")
    ss_papers = search_semantic_scholar(english_keyword, limit=limit, years_back=years_back)
    log_fn(f"[INFO] Semantic Scholar 결과: {len(ss_papers)}편")

    # 병합
    merged = merge_and_deduplicate(scholar_papers, openalex_papers, ss_papers)
    log_fn(f"[INFO] 병합 후 총 {len(merged)}편 (중복 제거 완료)")

    if not merged:
        raise RuntimeError("검색 결과가 없습니다.")

    # Zotero에 추가
    log_fn("[INFO] Zotero에 저장 중...")
    with sqlite3.connect(db_path) as conn:
        conn.execute("PRAGMA foreign_keys = ON")
        conn.execute("PRAGMA busy_timeout = 10000")
        cur = conn.cursor()
        cur.execute("BEGIN IMMEDIATE")
        try:
            library_id = get_user_library_id(cur)
            collection_id = get_or_create_collection(cur, library_id, keyword.strip())
            created = linked = 0
            for paper in merged:
                before = find_existing_item_id(cur, paper.title, paper.doi, library_id)
                item_id = insert_paper_as_journal_article(cur, library_id, paper)
                attach_item_to_collection(cur, collection_id, item_id)
                if before:
                    linked += 1
                else:
                    created += 1
            conn.commit()
        except Exception:
            conn.rollback()
            raise

    log_fn(f"[INFO] Zotero 추가: {created}편 신규, {linked}편 기존 연결")
    log_fn(f"[DONE] 컬렉션 '{keyword}' 업데이트 완료")

    # 분석 요청 JSON 저장 (proposal 있을 때)
    if proposal.strip():
        save_dir = output_path.parent if output_path else Path(".")
        json_path = save_analysis_request(keyword, proposal, save_dir, db_path=db_path)
        log_fn(f"[INFO] 분석 요청 파일 저장됨: {json_path.name}")
        log_fn("[INFO] GUI에서 'AI 분석' 버튼을 클릭하거나 Claude Code에서 파일을 열어주세요.")

    # docx 생성
    if output_path is not None:
        log_fn(f"[INFO] docx 생성 중: {output_path}")
        generate_summary_docx(merged, keyword, output_path, log_fn=log_fn)


# ── GUI ───────────────────────────────────────────────────────────────────────

def run_gui() -> None:
    import tkinter as tk
    from tkinter import scrolledtext, messagebox

    root = tk.Tk()
    root.title("논문 검색 & Zotero 추가 도구")
    root.resizable(False, False)

    pad = {"padx": 10, "pady": 6}

    # ── 입력 영역 ──────────────────────────────────────────────────
    frm = tk.Frame(root)
    frm.pack(fill="x", **pad)

    tk.Label(frm, text="키워드", width=8, anchor="w").grid(row=0, column=0, sticky="w", pady=4)
    entry_kw = tk.Entry(frm, width=40, font=("Malgun Gothic", 10))
    entry_kw.grid(row=0, column=1, columnspan=3, sticky="ew", pady=4)

    tk.Label(frm, text="논문 수", width=8, anchor="w").grid(row=1, column=0, sticky="w", pady=4)
    spin_limit = tk.Spinbox(frm, from_=1, to=20, width=5, font=("Malgun Gothic", 10))
    spin_limit.delete(0, "end")
    spin_limit.insert(0, "5")
    spin_limit.grid(row=1, column=1, sticky="w", pady=4)

    tk.Label(frm, text="최근").grid(row=1, column=2, sticky="w", padx=(16, 2))
    spin_years = tk.Spinbox(frm, from_=1, to=20, width=5, font=("Malgun Gothic", 10))
    spin_years.delete(0, "end")
    spin_years.insert(0, "5")
    spin_years.grid(row=1, column=3, sticky="w", pady=4)
    tk.Label(frm, text="년").grid(row=1, column=4, sticky="w")

    tk.Label(frm, text="내 연구\n제안", width=8, anchor="nw").grid(row=2, column=0, sticky="nw", pady=4)
    text_proposal = scrolledtext.ScrolledText(frm, width=40, height=4,
                                              font=("Malgun Gothic", 9), wrap="word")
    text_proposal.grid(row=2, column=1, columnspan=4, sticky="ew", pady=4)

    # ── 버튼 ───────────────────────────────────────────────────────
    frm_btn = tk.Frame(root)
    frm_btn.pack(pady=(0, 8))

    btn_search = tk.Button(frm_btn, text="  검색 시작  ", font=("Malgun Gothic", 11, "bold"),
                           bg="#2E74B5", fg="white", relief="flat", cursor="hand2")
    btn_search.pack(side="left", padx=4)

    btn_analyze = tk.Button(frm_btn, text="  AI 분석 (Claude)  ", font=("Malgun Gothic", 11, "bold"),
                            bg="#217346", fg="white", relief="flat", cursor="hand2",
                            state="disabled")
    btn_analyze.pack(side="left", padx=4)

    btn_nlm = tk.Button(frm_btn, text="  AI 분석 (NLM)  ", font=("Malgun Gothic", 11, "bold"),
                        bg="#7B2D8B", fg="white", relief="flat", cursor="hand2",
                        state="disabled")
    btn_nlm.pack(side="left", padx=4)

    last_json_path: list[Optional[Path]] = [None]
    last_keyword: list[str] = [""]
    last_db_path: list[Optional[Path]] = [None]

    # ── 로그 창 ────────────────────────────────────────────────────
    log_box = scrolledtext.ScrolledText(root, width=62, height=18,
                                        font=("Malgun Gothic", 9),
                                        state="disabled", bg="#F7F7F7")
    log_box.pack(fill="both", expand=True, padx=10, pady=(0, 10))

    def append_log(msg: str) -> None:
        log_box.configure(state="normal")
        log_box.insert("end", msg + "\n")
        log_box.see("end")
        log_box.configure(state="disabled")
        root.update_idletasks()

    def on_analyze() -> None:
        jp = last_json_path[0]
        if not jp or not jp.exists():
            messagebox.showwarning("분석 불가", "먼저 검색을 실행하고 연구 제안을 입력하세요.")
            return
        open_claude_terminal(jp, log_fn=append_log)

    btn_analyze.configure(command=on_analyze)

    def on_analyze_nlm() -> None:
        keyword = last_keyword[0]
        proposal_text = text_proposal.get("1.0", "end").strip()
        db = last_db_path[0]
        if not keyword or not proposal_text or not db:
            messagebox.showwarning("분석 불가", "먼저 검색을 실행하고 연구 제안을 입력하세요.")
            return
        btn_nlm.configure(state="disabled", text="  분석 중...  ")
        def nlm_worker() -> None:
            try:
                open_notebooklm_analysis(db, keyword, proposal_text, append_log)
            except Exception as exc:
                append_log(f"[ERROR] NLM 오류: {exc}")
            finally:
                root.after(0, lambda: btn_nlm.configure(state="normal",
                                                          text="  AI 분석 (NLM)  "))
        threading.Thread(target=nlm_worker, daemon=True).start()

    btn_nlm.configure(command=on_analyze_nlm)

    def on_search() -> None:
        keyword = entry_kw.get().strip()
        if not keyword:
            messagebox.showwarning("입력 오류", "키워드를 입력하세요.")
            return
        try:
            limit = int(spin_limit.get())
            years_back = int(spin_years.get())
        except ValueError:
            messagebox.showwarning("입력 오류", "논문 수와 연도는 숫자여야 합니다.")
            return

        safe_kw = _safe_filename(keyword)
        output_path = Path(f"{safe_kw}.docx")

        btn_search.configure(state="disabled", text="  검색 중...  ")
        btn_analyze.configure(state="disabled")
        btn_nlm.configure(state="disabled")
        log_box.configure(state="normal")
        log_box.delete("1.0", "end")
        log_box.configure(state="disabled")
        append_log(f"[START] 키워드: {keyword}  |  논문 수: {limit}  |  최근 {years_back}년")

        proposal = text_proposal.get("1.0", "end").strip()

        def worker() -> None:
            try:
                db_path = resolve_zotero_paths().db_path
                last_db_path[0] = db_path
                last_keyword[0] = keyword
                run_pipeline(db_path, keyword, limit, output_path,
                             years_back=years_back, log_fn=append_log,
                             proposal=proposal)
                append_log(f"\n[완료] 저장 파일: {output_path.resolve()}")
                # 분석 버튼 활성화
                if proposal:
                    jp = output_path.parent / f"{safe_kw}_analysis.json"
                    claude_ready = jp.exists()
                    if claude_ready:
                        last_json_path[0] = jp
                        root.after(0, lambda: btn_analyze.configure(state="normal"))
                    root.after(0, lambda: btn_nlm.configure(state="normal"))
                    active = "'AI 분석 (Claude)', 'AI 분석 (NLM)'" if claude_ready else "'AI 분석 (NLM)'"
                    append_log(f"[INFO] {active} 버튼이 활성화되었습니다.")
            except Exception as exc:
                append_log(f"[ERROR] {exc}")
            finally:
                root.after(0, lambda: btn_search.configure(state="normal", text="  검색 시작  "))

        threading.Thread(target=worker, daemon=True).start()

    btn_search.configure(command=on_search)
    entry_kw.bind("<Return>", lambda _: on_search())

    root.mainloop()


# ── CLI ───────────────────────────────────────────────────────────────────────

def run_cli() -> int:
    parser = argparse.ArgumentParser(
        description="키워드로 Google Scholar + OpenAlex + Semantic Scholar 논문을 검색해 Zotero에 추가하고 docx 요약을 생성합니다."
    )
    parser.add_argument("keyword", help="검색 키워드 (한글/영문 모두 가능)")
    parser.add_argument("--limit", type=int, default=5, help="소스별 최대 논문 수 (기본: 5)")
    parser.add_argument("--years", type=int, default=None, help="최근 N년 논문만 검색 (기본: 제한 없음)")
    parser.add_argument("--db-path", type=Path, default=None, help="zotero.sqlite 경로 (선택)")
    parser.add_argument("--output", type=Path, default=None, help="docx 저장 경로 (선택)")
    parser.add_argument("--no-docx", action="store_true", help="docx 생성 안 함")
    parser.add_argument("--proposal", type=str, default="", help="내 연구 제안 (AI 분석 활성화)")
    args = parser.parse_args()

    if args.limit < 1 or args.limit > 20:
        print("[ERROR] --limit는 1~20 사이여야 합니다.")
        return 1

    db_path = args.db_path or resolve_zotero_paths().db_path

    if args.no_docx:
        output_path = None
    elif args.output:
        output_path = args.output
    else:
        safe_kw = _safe_filename(args.keyword)
        output_path = Path(f"{safe_kw}.docx")

    try:
        run_pipeline(db_path, args.keyword, args.limit, output_path,
                     years_back=args.years, proposal=args.proposal)
    except sqlite3.OperationalError as exc:
        print("[ERROR] SQLite 오류:", exc)
        print("[HINT] Zotero를 종료한 뒤 다시 실행하세요.")
        return 1
    except urllib.error.HTTPError as exc:
        print("[ERROR] HTTP 오류:", exc)
        return 1
    except Exception as exc:
        print("[ERROR]", exc)
        return 1

    return 0


def main() -> int:
    if len(sys.argv) > 1:
        return run_cli()
    run_gui()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
