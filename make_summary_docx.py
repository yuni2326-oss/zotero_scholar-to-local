from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

FONT = 'Malgun Gothic'  # 맑은 고딕 영문명

def set_run_font(run, font_name=FONT):
    """한국어 포함 텍스트를 위해 ASCII/EastAsia 폰트 모두 설정"""
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

doc = Document()

# --- 기본 스타일 설정 ---
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(10)
# Normal 스타일에 eastAsia 폰트도 설정
from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement as _OE
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(_qn('w:rFonts'))
if rFonts is None:
    rFonts = _OE('w:rFonts')
    rPr.insert(0, rFonts)
rFonts.set(_qn('w:eastAsia'), FONT)

def set_heading(paragraph, text, level=1):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    if not paragraph.runs:
        run = paragraph.add_run(text)
    set_run_font(run)
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_section_heading(doc, text, level=2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run)
    if level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x20, 0x60, 0x40)
    return p

def add_paper(doc, num, title, authors, journal, doi, summary):
    # 논문 제목
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {title}")
    set_run_font(run)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 메타 정보
    meta = doc.add_paragraph()
    meta.paragraph_format.left_indent = Inches(0.2)
    r1 = meta.add_run("저자: ")
    r1.font.bold = True; r1.font.size = Pt(9); set_run_font(r1)
    r2 = meta.add_run(authors)
    r2.font.size = Pt(9); set_run_font(r2)
    meta.add_run("  |  ")
    r3 = meta.add_run("저널: ")
    r3.font.bold = True; r3.font.size = Pt(9); set_run_font(r3)
    r4 = meta.add_run(journal)
    r4.font.size = Pt(9); set_run_font(r4)
    meta.add_run("  |  ")
    r5 = meta.add_run("DOI: ")
    r5.font.bold = True; r5.font.size = Pt(9); set_run_font(r5)
    r6 = meta.add_run(doi)
    r6.font.size = Pt(9); r6.font.color.rgb = RGBColor(0x00, 0x70, 0xC0); set_run_font(r6)

    # 요약
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.2)
    rl = p2.add_run("요약: ")
    rl.font.bold = True; rl.font.size = Pt(10); set_run_font(rl)
    rs = p2.add_run(summary)
    rs.font.size = Pt(10); set_run_font(rs)

    doc.add_paragraph()  # 간격

# ========== 문서 시작 ==========

add_title(doc, "나노플라스틱 관련 논문 요약 보고서")

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_date.add_run(f"작성일: {datetime.date.today().strftime('%Y년 %m월 %d일')}  |  출처: Zotero 라이브러리")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
set_run_font(r)

doc.add_paragraph()
doc.add_paragraph()

# ========== 1. 검출 방법론 ==========
add_section_heading(doc, "1. 검출 방법론 (Detection Methods)", level=2)

add_paper(doc, 1,
    "Methodological Controversies in Microplastic Detection: A Critical Analysis of Pyrolysis-GC-MS False Positives in Human Tissue Studies",
    "Zen Revista",
    "Zenodo (CERN), 2026",
    "10.5281/zenodo.18518443",
    "인체 장기에서 마이크로플라스틱 오염을 보고한 고프로파일 연구들에 대한 방법론적 논란을 분석. "
    "특히 Pyrolysis-GC-MS 기법의 한계를 지적하며, 지방 함량이 높은 뇌 조직(약 60% 지방)이 폴리에틸렌과 유사한 열분해 산물을 생성해 "
    "위양성(false positive)을 유발할 수 있음을 제시. 공식적으로 이의 제기된 7개 연구와 행렬 간섭을 충분히 고려하지 않은 18개 연구를 추가 식별. "
    "인체 건강 위험에 대한 결론을 내리기 위해서는 검증 프로토콜, 오염 관리, 표준화된 방법론이 시급히 필요하다고 강조."
)

add_paper(doc, 2,
    "Recent Progress and Technological Advancements for Detection of Micro/Nano-Plastics in the Environment",
    "Ashok Kumar; Monika Nehra; Nitin Kumar Singhal; Sandeep Kumar",
    "Advances in Colloid and Interface Science, 2026",
    "10.1016/j.cis.2026.103817",
    "수계·육상·대기·식품 환경에서 마이크로/나노플라스틱(MNPs) 검출을 위한 최신 분석 전략을 종합 검토. "
    "이미징 기반 기법, 분광법, 전기화학적 센서 플랫폼, 인공지능(AI) 통합 등을 분류하여 평가. "
    "콜로이드 계면과학적 관점에서 MNPs의 물리화학적 특성이 센서 성능에 미치는 영향을 분석하고, "
    "하이퍼스펙트럴 이미징 등 원격 모니터링 플랫폼의 가능성을 논의. "
    "다중 검출 방식의 통합을 통해 실시간, 확장 가능한 MNPs 모니터링 실현이 가능하다고 전망."
)

add_paper(doc, 3,
    "Development of a Dye-Binding Method for Nanoplastics Detection in Snow Samples Using Capillary Electrophoresis with Laser-Induced Fluorescence Detection",
    "Edward P.C. Lai; Amos Onomhante; Apollinaire Tsopmo; Farah Hosseinian",
    "ACS Omega, 2026",
    "10.1021/acsomega.5c13510",
    "눈 시료 내 나노플라스틱 검출을 위한 새로운 형광 염료-모세관 전기영동-레이저 유도 형광(CE-LIF) 방법 개발. "
    "형광 유기 염료를 나노플라스틱에 결합시킨 후 원심분리로 분리하고, 15분 내에 최소 0.1 μg의 나노플라스틱을 신속 정량 가능. "
    "주성분 분석(PCA)으로 시료 간 나노플라스틱 유형·표면 화학·입자 크기 분포의 변동성을 확인. "
    "최소한의 시료 처리로 높은 분석 처리량을 달성하는 실용적 방법으로 평가됨."
)

add_paper(doc, 4,
    "Detection and Fate of Microplastics and Nanoplastics and Technologies for Their Removal",
    "Qiuping Zhang; Qi Wang; JiFei Xu; Jianguo Liu",
    "Molecules, 2026",
    "10.3390/molecules31040613",
    "연간 4억 6천만 메트릭 톤 이상 생산되는 플라스틱의 분해 산물인 MP/NP의 검출·환경 거동·제거 기술을 체계적으로 검토. "
    "검출 기법으로 현미경, 질량분석, 유세포분석, 크로마토그래피, 분광법(FT-IR, 라만)을 다루며, "
    "환경 내 장기 잔류성, 생물 농축, 먹이사슬 전달, 오염물 운반체 역할을 분석. "
    "물리적(막 여과, 흡착), 화학적(응집, 고급 산화), 생물학적(바이오숯, 미생물 분해) 제거 방법의 장단점을 비교 평가."
)

doc.add_page_break()

# ========== 2. 인체 건강 영향 ==========
add_section_heading(doc, "2. 인체 건강 영향 (Human Health Effects)", level=2)

add_paper(doc, 5,
    "Why Detecting Nanoplastics in Humans Matters: Exposure Routes, Biological Evidence, and Potential Health Implications",
    "Pushpender Kumar",
    "2026",
    "10.62830/mmj3-01-5a",
    "1 μm 미만의 나노플라스틱이 혈액, 폐 조직, 태반, 대변, 소변, 혈관 조직, 뇌 등 인체 다양한 생물학적 기질에서 검출됨. "
    "섭취와 흡입이 주요 노출 경로이며, 모체-태아 간 이동도 우려 대상. "
    "실험 연구에서 산화 스트레스, 염증, 내분비 교란, 신경생물학적 영향이 확인되나 인체에서의 직접적 인과관계는 아직 제한적. "
    "Pyrolysis GC-MS, 라만 분광법, 단일 입자 ICP-MS 등의 분석 기법 표준화가 임상 적용을 위해 필수적임을 강조."
)

add_paper(doc, 6,
    "Human Exposure to Micro- and Nanoplastics: Routes, Accumulation, and Health Implications",
    "Min Young Oh; Gyeong Bae Song; Bokyung Kim; Jeein Hong; Minseon Ju; Sungguan Hong",
    "Molecular & Cellular Toxicology, 2026",
    "10.1007/s13273-026-00610-0",
    "마이크로 및 나노플라스틱에 대한 인체 노출 경로, 체내 축적 메커니즘, 건강 영향을 종합적으로 분석한 리뷰 논문. "
    "음식 섭취, 음용수, 공기 흡입 등 다양한 노출 경로를 통해 인체 조직에 축적되며, "
    "독성학적·세포 수준에서의 영향과 잠재적 건강 위험을 검토함."
)

add_paper(doc, 7,
    "Microplastics as an Emerging Human Health Risk: Mechanisms, Exposure, and Clinical Evidence",
    "Dr. Asif Rasheed; Hadiqa Bushra Awesi; Farzeen Ajmal Younus",
    "Zenodo (CERN), 2026",
    "10.5281/zenodo.18523585",
    "5 mm 미만의 마이크로플라스틱이 혈액, 태반, 폐 조직, 심혈관 조직 등 인체 생물학적 시료에서 광범위하게 검출. "
    "생물막 통과 후 전신 분포 및 장기별 축적이 확인되며, 산화 스트레스·염증·상피 장벽 붕괴·미토콘드리아 기능 이상·혈전 형성 등이 주요 독성 기전. "
    "심혈관 질환, 호흡기 장애, 염증 질환과의 임상적 연관성이 보고되나 인과관계는 추가 연구 필요. "
    "예방 전략으로 플라스틱 사용 감소, 식이 습관 개선, 수처리 시스템 향상을 제시."
)

add_paper(doc, 8,
    "Micro-Nanoplastics in the Central Nervous System: Evidence, Mechanisms and Perspectives",
    "Zhexun Pei; Hang Zhong; Xiaoqing Li; Shaowei Guo; Tingting Wang",
    "Toxicology, 2026",
    "10.1016/j.tox.2026.154424",
    "나노플라스틱이 혈액-뇌 장벽(BBB)을 통과하여 신경독성을 유발한다는 증거를 종합 검토. "
    "산화 스트레스, 신경염증, 미토콘드리아 기능 장애, 신경전달물질 교란이 주요 메커니즘으로 확인. "
    "설치류 모델에서 신경퇴행, 신경발달 장애, 정신질환 관련 행동·병리학적 결손이 관찰됨. "
    "인체 중추신경계 조직 및 체액에서 나노플라스틱 존재가 확인되나, 방법론적 한계로 인과 추론에 제약. "
    "BBB 무결성, 인지 기능, 뇌졸중 중증도와 나노플라스틱 수준 간의 임상적 상관관계가 보고됨."
)

add_paper(doc, 9,
    'Correspondence Regarding "Inflammatory Effects of Microplastics and Nanoplastics on Nasal Airway Epithelial Cells"',
    "Hyun Jin Min",
    "International Forum of Allergy & Rhinology, 2026",
    "10.1002/alr.70124",
    "비강 기도 상피세포에 대한 마이크로/나노플라스틱의 염증 효과를 다룬 Kahan 등의 연구에 대한 서신. "
    "폴리스티렌 입자 크기와 표면 전하에 따른 세포 반응 차이를 분석한 해당 연구를 평가하며, "
    "적절한 대조군 설정, 기저측 노출 실험 필요성, 노출 시점 선택 기준 등 방법론적 개선 사항을 제안. "
    "흡입된 나노플라스틱이 상기도 염증을 유발하고 호흡기 건강에 위험을 초래할 수 있음을 시사."
)

add_paper(doc, 10,
    "Evaluating Cellular Effects of PET Microplastics in 2D/3D Models: Methodological Considerations of Reagent Interference",
    "Lisa Aichinger 외 7인",
    "2026",
    "10.21203/rs.3.rs-8672251/v1",
    "폴리에틸렌 테레프탈레이트(PET) 마이크로플라스틱(직경 1.24 μm)의 세포 독성을 2D 단층 배양 및 3D 다세포 구형체 모델에서 평가. "
    "2.5–5 μg/ml 농도에서 세포 생존율이 농도·세포 유형 의존적으로 감소하며, 세포 내 흡수율은 최대 16%. "
    "단기 노출(10–120분)에서 농도 의존적 활성산소종(ROS) 생성이 유도됨. "
    "형광 기반 분석에서 플라스틱 흡착, 광산란, 계면활성제 효과 등으로 인한 실험 인공물 가능성을 경고. "
    "O-PTIR 분광법이 비표지 PET 입자의 생체 내 분포 분석에 유용함을 제시."
)

doc.add_page_break()

# ========== 3. 수생 생태계 오염 ==========
add_section_heading(doc, "3. 수생 생태계 오염 (Aquatic Ecosystem Contamination)", level=2)

add_paper(doc, 11,
    "From Rivers to Humans: Evolving Policies and Health Risks of Microplastics Pollution in Freshwater Fish",
    "Prachi Kumari; Akshay Gautam; Prateek Kumar Tiwari; Satyendra Katara",
    "Journal of Advances in Biology & Biotechnology, 2026",
    "10.9734/jabb/2026/v29i23703",
    "농업 활동, 도시 유출수, 하수 방류, 대기 침적 등으로 담수 생태계에 광범위하게 유입된 마이크로/나노플라스틱을 분석. "
    "민물고기는 섭식 및 아가미 흡수를 통해 지속적으로 노출되며, 주로 위장관에 집중되나 식용 근육 조직에서의 존재는 연구마다 상이. "
    "2017–2025년 인도 및 글로벌 연구를 검토하여 발생 패턴, 독성 영향, 인체 노출 경로, 분석 과제, 정책 대응을 종합. "
    "인체 혈액, 태반, 동맥 조직에서 마이크로플라스틱 검출 보고가 공중 보건 우려를 증폭. EU·인도 규제 조치의 한계와 향후 방향을 논의."
)

add_paper(doc, 12,
    "Microplastic Pollution in Aquatic Ecosystems: Environmental Behaviour, Biological Impacts, and Public Health Implications",
    "K. Nikhil 외 7인",
    "Uttar Pradesh Journal of Zoology, 2026",
    "10.56557/upjoz/2026/v47i45529",
    "담수 및 해양 시스템에서 마이크로/나노플라스틱의 주요 발생원, 이동 경로, 물리화학적 특성, 환경 거동(운반·파편화·생물오손·침강)을 종합 검토. "
    "영양 단계별 생태적 영향과 인체 노출 경로(해산물, 음용수, 흡입)를 분석하며, "
    "염증·산화 스트레스·내분비 교란 등 독성 기전을 논의. "
    "FTIR/라만 분광법, 열분해-GC/MS 등 분석 검출 기법의 현황과 한계를 검토하고, "
    "규제·완화 접근법 및 향후 연구 우선순위(표준화 정의, 방법론 비교, 정량적 위험 평가 프레임워크)를 제시."
)

add_paper(doc, 13,
    "Ecotoxicological Assessment of Emerging Pollutants and Diatom Vulnerability in Aquatic Ecosystems",
    "Anuradha Yadav; Arpita Srivastava; Jyoti Verma",
    "Discover Water, 2026",
    "10.1007/s43832-025-00339-2",
    "수생 생태계 건강의 민감한 지표인 규조류(diatom)에 대한 신흥 오염물질(농약, PFAS, 마이크로/나노플라스틱, 의약품 등)의 영향을 분석. "
    "150개 이상의 생태독성 연구와 Umiam 저수지(메갈라야) 현장 데이터를 종합하여 생물다양성 감소와 영양 순환 방해를 확인. "
    "기후 스트레스가 신흥 오염물질의 영향을 증폭시키며, 장기 저용량 복합 노출이 규조류 광합성·규소질 껍질 생산·유전자 발현을 교란. "
    "O3/UV 고급 산화 기술, 조류-박테리아 반응기, 인공 습지를 부분적 복원 방법으로 제시. "
    "실시간 기계학습 모니터링, 오믹스 기반 바이오마커, 스마트폰 현미경 등 미래 기술을 전망."
)

add_paper(doc, 14,
    "Nanoplastics in Biological Systems: What Laboratory Mechanisms Reveal about Real-World Toxicity",
    "Fernan Arellano; Cris Gel Loui A. Arcadio; Ya-Ting Chen 외",
    "Journal of Hazardous Materials, 2026",
    "10.1016/j.jhazmat.2026.141464",
    "수계·대기·식품 시스템에서 나노플라스틱이 증가하고 있으나, 고농도 실험실 연구와 실제 환경 저농도 만성 노출 간의 괴리가 문제. "
    "세포 흡수, 리소좀 파열, 미토콘드리아 기능 장애, 산화 스트레스, 염증 활성화 등 핵심 독성 경로가 환경 관련 농도에서도 유지됨을 제시. "
    "환경적 노화, 에코코로나 형성, 공동 오염물질 부하가 낮은 농도에서도 나노플라스틱의 반응성을 증폭. "
    "AI/머신러닝이 나노플라스틱 검출, 노화 입자 특성화, 오염물질 상호작용 예측에 혁신적 전환을 이루고 있음. "
    "농도 기반 평가에서 확률론적·메커니즘 기반 모델로의 전환을 촉구."
)

doc.add_page_break()

# ========== 4. 노출 및 축적 모델링 ==========
add_section_heading(doc, "4. 노출 및 축적 모델링 (Exposure & Accumulation Modeling)", level=2)

add_paper(doc, 15,
    "Quantitative Modeling of Nanoplastic Accumulation from Single-Use Water Bottles: Exposure Scenarios and Tissue Concentration Estimates",
    "Sarah Sajedi; Chunjiang An",
    "Environmental Systems Research, 2026",
    "10.1186/s40068-026-00459-1",
    "일회용 플라스틱 물병에서 나노플라스틱의 장기 섭취로 인한 인체 조직 축적을 정량적으로 모델링. "
    "HEASI Plastic Model을 이용하여 1.10×10⁵~1.0×10¹¹ particles/L 범위의 3가지 노출 시나리오를 설정. "
    "전신 조직 내 나노플라스틱 축적은 소비 습관과 담즙 배설률에 따라 0.00084~226.68 μg/L로 광범위하게 변동. "
    "폴리아미드와 폴리스티렌이 조직 내 주요 나노플라스틱 유형으로 모델링됨. "
    "고소비·고노출 시나리오에서 용량 의존적 축적 경향이 확인되며, 규제 기준 설정과 일회용 플라스틱 사용 감소를 촉구."
)

# ========== 맺음말 ==========
doc.add_page_break()
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_end = p_end.add_run("── 요약 보고서 끝 ──")
r_end.font.size = Pt(9)
r_end.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
set_run_font(r_end)

note = doc.add_paragraph()
note.paragraph_format.left_indent = Inches(0.2)
rn = note.add_run(
    "※ 본 보고서는 Zotero 라이브러리 내 '나노 플라스틱 검출(kw:나노 플라스틱 검출)' 태그 논문 15편의 초록을 기반으로 작성되었습니다.\n"
    "   중복 항목(IA8BX7AU, LGEWN954, 4R2PSR7D, QPUWYMMG, US8WB3PU, A3TY53CT 일부)은 제외하거나 대표 항목으로 통합하였습니다."
)
rn.font.size = Pt(8)
rn.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
set_run_font(rn)

doc.save("e:/My project/Zotero_test/나노플라스틱_논문요약.docx")
print("저장 완료: 나노플라스틱_논문요약.docx")
